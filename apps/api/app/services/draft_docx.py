"""초안 DOCX 변환과 저장 (PRD §7 F4).

`content_json` 을 그대로 문서로 옮긴다. 여기서 문구를 새로 만들지 않는다.

- **운영명세서(sow)**: 머리글 1행 + 항목 101행짜리 표.
- **정책(policy)**: 제목 + 조항별 소제목·본문 문단.

한글이 네모로 깨지지 않게 본문·제목 스타일에 한글 글꼴(맑은 고딕)을 명시한다.
글꼴이 없는 환경에서는 워드가 알아서 대체 글꼴을 쓴다.
"""

import io
import uuid
from typing import Any

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.ns import qn
from docx.shared import Pt

from app.models import DraftKind, Project
from app.services.draft_common import NEEDS_REVIEW
from app.services.storage import get_storage

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 한글 글꼴. 설치돼 있지 않은 환경에서는 워드가 대체 글꼴로 렌더한다.
KOREAN_FONT_NAME = "맑은 고딕"
# 글꼴을 명시할 스타일. 표·제목까지 같은 글꼴로 맞춘다.
_FONT_STYLE_NAMES = ("Normal", "Title", "Heading 1", "Heading 2", "Table Grid")

# 운영명세서 표 열 머리글(PRD §7 F4).
SOW_COLUMNS = ("항목 코드", "항목명", "운영 현황", "관련 문서·증적", "담당 부서", "비고")
SOW_TABLE_STYLE = "Table Grid"

_BODY_FONT_SIZE = Pt(9)

_KIND_TITLES = {
    DraftKind.SOW: "정보보호 관리체계 운영명세서",
    DraftKind.POLICY: "정보보호 정책",
}

_DRAFT_NOTICE = (
    "이 문서는 CertPilot 이 생성한 초안이다. 심사원 검수·승인 전에는 확정본이 아니며, "
    f"{NEEDS_REVIEW} 로 표시된 칸은 담당자가 직접 확인해 채워야 한다."
)


class DraftRenderError(RuntimeError):
    """초안을 DOCX 로 옮길 수 없다(content_json 형식이 어긋남)."""


def _apply_korean_font(document: DocxDocumentType) -> None:
    """본문·제목·표 스타일에 한글 글꼴을 명시한다.

    워드는 동아시아 문자에 `w:eastAsia` 글꼴을 따로 쓴다. 이 값을 지정하지 않으면
    한글이 기본 라틴 글꼴로 렌더돼 깨져 보인다.
    """
    for name in _FONT_STYLE_NAMES:
        if name not in document.styles:  # pragma: no cover - 기본 템플릿에는 모두 있다
            continue
        style = document.styles[name]
        style.font.name = KOREAN_FONT_NAME
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attribute), KOREAN_FONT_NAME)
    document.styles["Normal"].font.size = _BODY_FONT_SIZE


def _add_meta(document: DocxDocumentType, project: Project, kind: DraftKind, version: int) -> None:
    """문서 머리에 프로젝트·버전 정보와 초안 고지를 넣는다."""
    document.add_heading(f"{project.name} {_KIND_TITLES[kind]}", level=0)
    meta = document.add_paragraph()
    meta.add_run(
        f"인증 종류: {project.cert_type.value} · 버전: v{version} · "
        f"인증범위: {(project.scope_text or NEEDS_REVIEW).strip()}"
    )
    notice = document.add_paragraph()
    notice.add_run(_DRAFT_NOTICE).italic = True


def _rows_of(content: dict[str, Any]) -> list[dict[str, Any]]:
    """`content_json` 에서 운영명세서 행을 꺼낸다."""
    rows = content.get("rows")
    if not isinstance(rows, list):
        raise DraftRenderError("운영명세서 content_json 에 rows 배열이 없다")
    return [row for row in rows if isinstance(row, dict)]


def _sections_of(content: dict[str, Any]) -> list[dict[str, Any]]:
    """`content_json` 에서 정책 조항을 꺼낸다."""
    sections = content.get("sections")
    if not isinstance(sections, list):
        raise DraftRenderError("정책 초안 content_json 에 sections 배열이 없다")
    return [section for section in sections if isinstance(section, dict)]


def render_sow_docx(project: Project, content: dict[str, Any], version: int) -> bytes:
    """운영명세서 DOCX 를 만든다. 표는 머리글 1행 + 항목 수만큼의 행이다."""
    rows = _rows_of(content)

    document = DocxDocument()
    _apply_korean_font(document)
    _add_meta(document, project, DraftKind.SOW, version)

    stats = content.get("stats") or {}
    document.add_paragraph(
        f"총 {stats.get('total', len(rows))}개 항목 · "
        f"확인이 필요한 칸 {stats.get('needs_review', 0)}개"
    )

    table = document.add_table(rows=1, cols=len(SOW_COLUMNS))
    table.style = SOW_TABLE_STYLE
    for cell, header in zip(table.rows[0].cells, SOW_COLUMNS, strict=True):
        cell.text = ""
        cell.paragraphs[0].add_run(header).bold = True

    for row in rows:
        cells = table.add_row().cells
        related = row.get("related_refs") or []
        cells[0].text = str(row.get("criterion_code", ""))
        cells[1].text = str(row.get("criterion_title", ""))
        cells[2].text = str(row.get("operation_status", ""))
        cells[3].text = "\n".join(str(item) for item in related)
        cells[4].text = str(row.get("owner_dept", NEEDS_REVIEW))
        cells[5].text = str(row.get("note", ""))

    return _to_bytes(document)


def render_policy_docx(project: Project, content: dict[str, Any], version: int) -> bytes:
    """정책 초안 DOCX 를 만든다. 조항마다 소제목 + 본문 문단."""
    sections = _sections_of(content)

    document = DocxDocument()
    _apply_korean_font(document)
    _add_meta(document, project, DraftKind.POLICY, version)

    for section in sections:
        document.add_heading(str(section.get("heading", "")), level=1)
        for line in str(section.get("body", "")).split("\n"):
            if line.strip():
                document.add_paragraph(line.strip())

    return _to_bytes(document)


def render_draft_docx(
    kind: DraftKind, project: Project, content: dict[str, Any], version: int
) -> bytes:
    """초안 종류에 맞는 DOCX 를 만든다."""
    if kind is DraftKind.SOW:
        return render_sow_docx(project, content, version)
    return render_policy_docx(project, content, version)


def _to_bytes(document: DocxDocumentType) -> bytes:
    """메모리에서 DOCX 바이트로 직렬화한다."""
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def draft_filename(kind: DraftKind, version: int) -> str:
    """다운로드 파일명. 한글 파일명은 헤더 인코딩 문제가 있어 ASCII 로 만든다."""
    return f"certpilot-{kind.value}-v{version}.docx"


def draft_docx_key(
    *, org_id: uuid.UUID, project_id: uuid.UUID, draft_id: uuid.UUID, kind: DraftKind, version: int
) -> str:
    """S3 오브젝트 키. 문서 업로드와 같은 org/project 접두사를 쓴다."""
    return (
        f"orgs/{org_id}/projects/{project_id}/drafts/{draft_id}/"
        f"{draft_filename(kind, version)}"
    )


def store_draft_docx(key: str, payload: bytes) -> None:
    """DOCX 를 오브젝트 스토리지에 올린다. 실패는 `StorageError` 로 올라간다."""
    get_storage().put_object(key, payload, DOCX_MEDIA_TYPE)


def load_draft_docx(key: str) -> bytes:
    """저장된 DOCX 를 읽는다. 실패는 `StorageError` 로 올라간다."""
    return get_storage().get_object(key)
