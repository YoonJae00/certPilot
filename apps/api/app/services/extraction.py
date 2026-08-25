"""업로드 문서에서 텍스트를 뽑아낸다(pdf / docx / xlsx / md).

여기서는 마스킹을 하지 않는다. 마스킹은 `app.services.masking` 이 청킹 직전에
한 번만 적용한다(단일 지점 유지).
"""

import io
import logging
from dataclasses import dataclass

import openpyxl
import pdfplumber
from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

logger = logging.getLogger(__name__)

# 확장자 → MIME. 업로드 허용 목록이자 파서 선택 키다.
SUPPORTED_EXTENSIONS: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "md": "text/markdown",
}


class ExtractionError(RuntimeError):
    """텍스트 추출 실패. 메시지는 사용자에게 그대로 보여줄 한국어 사유다."""


@dataclass(frozen=True)
class ExtractedPage:
    """추출된 페이지(또는 시트) 1개."""

    page: int
    text: str


@dataclass(frozen=True)
class ExtractedDocument:
    """문서 전체 추출 결과."""

    pages: list[ExtractedPage]

    @property
    def page_count(self) -> int:
        """페이지(시트) 수."""
        return len(self.pages)

    @property
    def text(self) -> str:
        """전체 텍스트."""
        return "\n\n".join(page.text for page in self.pages)


def _extract_pdf_with_pdfplumber(data: bytes) -> list[ExtractedPage]:
    """pdfplumber 로 페이지별 텍스트를 뽑는다."""
    pages: list[ExtractedPage] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            pages.append(ExtractedPage(page=index, text=page.extract_text() or ""))
    return pages


def _extract_pdf_with_pypdf(data: bytes) -> list[ExtractedPage]:
    """pypdf 폴백. pdfplumber 가 못 여는 PDF 를 한 번 더 시도한다."""
    reader = PdfReader(io.BytesIO(data))
    return [
        ExtractedPage(page=index, text=page.extract_text() or "")
        for index, page in enumerate(reader.pages, start=1)
    ]


def _is_encrypted_pdf(data: bytes) -> bool:
    """암호로 보호된 PDF 인지 확인한다."""
    try:
        return bool(PdfReader(io.BytesIO(data)).is_encrypted)
    except (PdfReadError, ValueError, OSError):
        # 헤더조차 못 읽는 경우다. 암호 문제가 아니므로 여기서는 False.
        return False


def extract_pdf(data: bytes) -> list[ExtractedPage]:
    """PDF 텍스트 추출. pdfplumber 우선, 실패하면 pypdf 로 폴백한다."""
    if _is_encrypted_pdf(data):
        raise ExtractionError("암호로 보호된 PDF 라 텍스트를 추출할 수 없다")

    try:
        return _extract_pdf_with_pdfplumber(data)
    except Exception as primary_error:  # noqa: BLE001 - 폴백 후 원인을 다시 올린다
        logger.warning("pdfplumber 추출 실패, pypdf 로 폴백한다: %s", primary_error)
        try:
            return _extract_pdf_with_pypdf(data)
        except Exception as fallback_error:
            raise ExtractionError(
                f"PDF 텍스트 추출 실패: {fallback_error}"
            ) from fallback_error


def extract_docx(data: bytes) -> list[ExtractedPage]:
    """DOCX 문단과 표를 텍스트로 편다.

    DOCX 에는 페이지 개념이 없다(렌더링 시점에 정해진다). 문서 전체를 1페이지로 본다.
    """
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as error:
        raise ExtractionError(f"DOCX 열기 실패: {error}") from error

    lines: list[str] = [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return [ExtractedPage(page=1, text="\n".join(lines))]


def extract_xlsx(data: bytes) -> list[ExtractedPage]:
    """XLSX 를 시트 단위로 편다. 시트 1개가 페이지 1개다."""
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except Exception as error:
        raise ExtractionError(f"XLSX 열기 실패: {error}") from error

    try:
        pages: list[ExtractedPage] = []
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = [f"[시트] {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value).strip() for value in row]
                if any(cells):
                    lines.append(" | ".join(cells))
            pages.append(ExtractedPage(page=index, text="\n".join(lines)))
        return pages
    finally:
        workbook.close()


def extract_markdown(data: bytes) -> list[ExtractedPage]:
    """마크다운은 그대로 텍스트로 쓴다."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ExtractionError("UTF-8 로 읽을 수 없는 텍스트 파일이다") from error
    return [ExtractedPage(page=1, text=text)]


def extension_of(filename: str) -> str:
    """파일명에서 소문자 확장자를 뽑는다."""
    _, _, extension = filename.rpartition(".")
    return extension.lower()


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    """확장자에 맞는 파서로 텍스트를 추출한다."""
    extension = extension_of(filename)
    if extension == "pdf":
        pages = extract_pdf(data)
    elif extension == "docx":
        pages = extract_docx(data)
    elif extension == "xlsx":
        pages = extract_xlsx(data)
    elif extension == "md":
        pages = extract_markdown(data)
    else:
        raise ExtractionError(f"지원하지 않는 파일 형식이다: .{extension}")

    if not any(page.text.strip() for page in pages):
        raise ExtractionError("문서에서 텍스트를 한 글자도 추출하지 못했다")
    return ExtractedDocument(pages=pages)
