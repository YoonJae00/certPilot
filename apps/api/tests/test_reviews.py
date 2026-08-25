"""검수 워크플로 테스트 (PRD §7 F6).

AC: §3 권한 매트릭스 그대로(조직 사용자는 검수 API 403, 심사원은 조직 API 403),
큐 조회·배정·편집(DOCX 재생성)·승인 후 다운로드 개방·반려 알림·중복 결정 409·감사 로그.

운영명세서 초안을 API 로 만들려면 모의심사를 먼저 돌려야 해서 느리다. 검수 로직 자체는
초안 본문 구조에만 의존하므로, 행 편집 검증은 작은 초안을 DB 에 직접 넣어 확인한다.
승인·반려·다운로드처럼 실제 파일이 필요한 경로는 API 로 만든 정책 초안을 쓴다.
"""

import io
import uuid

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import (
    Alert,
    AlertType,
    AuditLog,
    Draft,
    DraftKind,
    DraftStatus,
    ReviewTask,
    ReviewTaskStatus,
    UserRole,
)
from app.services.draft_common import NEEDS_REVIEW
from app.services.draft_docx import draft_docx_key
from app.services.review import ensure_review_task
from tests.conftest import login, make_user

REVIEWER_EMAIL = "reviewer@example.com"
OTHER_REVIEWER_EMAIL = "reviewer2@example.com"


# ---------------------------------------------------------------------------
# 픽스처·헬퍼
# ---------------------------------------------------------------------------


def task_of(db, draft_id) -> ReviewTask:
    """초안에 딸린 검수 과제를 읽는다."""
    db.expire_all()
    return db.execute(
        select(ReviewTask).where(ReviewTask.draft_id == uuid.UUID(str(draft_id)))
    ).scalar_one()


def load_draft(db, draft_id) -> Draft:
    """DB 에서 초안을 다시 읽는다(API 와 다른 세션이다)."""
    db.expire_all()
    return db.execute(select(Draft).where(Draft.id == uuid.UUID(str(draft_id)))).scalar_one()


@pytest.fixture
def policy_draft(client, db, tenants, storage):
    """조직 관리자가 만든 정책 초안 1건. 생성 직후 검수 대기 상태다."""
    login(client, "admin-a@example.com")
    response = client.post(
        f"/projects/{tenants['project_a'].id}/drafts", json={"kind": "policy"}
    )
    assert response.status_code == 201, response.text
    return response.json()


@pytest.fixture
def sow_draft(db, tenants, storage):
    """운영명세서 초안 3행짜리 축소판. 모의심사 없이 편집 경로만 검증한다."""
    project = tenants["project_a"]
    rows = [
        {
            "criterion_code": "1.1.1",
            "criterion_title": "경영진의 참여",
            "section": "1.1 관리체계 기반 마련",
            "operation_status": f"{NEEDS_REVIEW} 근거를 찾지 못해 판단할 수 없다.",
            "related_refs": ["정보보호정책 v2.1 p.3"],
            "owner_dept": NEEDS_REVIEW,
            "note": "",
        },
        {
            "criterion_code": "1.1.2",
            "criterion_title": "최고책임자의 지정",
            "section": "1.1 관리체계 기반 마련",
            "operation_status": "CISO 를 지정해 운영한다.",
            "related_refs": ["정보보호정책 v2.1 p.7"],
            "owner_dept": NEEDS_REVIEW,
            "note": "",
        },
        {
            "criterion_code": "2.1.1",
            "criterion_title": "정책의 유지관리",
            "section": "2.1 정책, 조직, 자산 관리",
            "operation_status": "연 1회 정책을 개정한다.",
            "related_refs": ["정보보호정책 v2.1 p.9"],
            "owner_dept": NEEDS_REVIEW,
            "note": "",
        },
    ]
    draft = Draft(
        project_id=project.id,
        kind=DraftKind.SOW,
        version=1,
        status=DraftStatus.IN_REVIEW,
        content_json={
            "rows": rows,
            "stats": {"total": 3, "needs_review": 4, "needs_review_rows": 3},
        },
        created_by=tenants["admin_a"].id,
    )
    draft.docx_s3_key = None
    db.add(draft)
    db.commit()
    db.refresh(draft)

    ensure_review_task(db, draft)
    db.commit()
    return draft


@pytest.fixture
def other_reviewer(db):
    """같은 큐를 보는 두 번째 심사원."""
    return make_user(db, email=OTHER_REVIEWER_EMAIL, role=UserRole.REVIEWER)


def open_task(client, task_id):
    """검수 과제 상세 조회(= 미배정이면 배정)."""
    return client.get(f"/reviews/{task_id}")


# ---------------------------------------------------------------------------
# 초안 생성 → 검수 과제
# ---------------------------------------------------------------------------


def test_draft_creation_enqueues_unassigned_task(client, db, policy_draft):
    """초안을 만들면 미배정 대기 과제가 하나 생긴다."""
    task = task_of(db, policy_draft["id"])
    assert task.status is ReviewTaskStatus.PENDING
    # 담당 심사원은 큐에서 잡을 때 정해진다.
    assert task.reviewer_id is None
    assert task.decided_at is None


def test_pending_task_is_not_duplicated(db, sow_draft):
    """같은 초안에 대기 과제를 두 개 만들지 않는다."""
    ensure_review_task(db, sow_draft)
    db.commit()

    tasks = list(
        db.execute(select(ReviewTask).where(ReviewTask.draft_id == sow_draft.id)).scalars()
    )
    assert len(tasks) == 1


# ---------------------------------------------------------------------------
# 권한 매트릭스 (PRD §3)
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("email", ["admin-a@example.com", "member-a@example.com"])
def test_org_users_cannot_use_review_api(client, db, policy_draft, email):
    """조직 관리자·팀원은 검수 API 에 접근할 수 없다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, email)

    assert client.get("/reviews/queue").status_code == 403
    assert open_task(client, task_id).status_code == 403
    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 403
    assert client.post(f"/reviews/{task_id}/return", json={"comment": "보완"}).status_code == 403
    assert (
        client.patch(
            f"/reviews/{task_id}/content", json={"sections": [{"section_index": 0, "body": "x"}]}
        ).status_code
        == 403
    )


def test_reviewer_is_still_blocked_from_org_api(client, db, policy_draft, tenants):
    """심사원은 검수 과제를 통해서만 접근한다. 조직 스코프 API 는 그대로 403 이다."""
    project_id = tenants["project_a"].id
    login(client, REVIEWER_EMAIL)

    assert client.get("/projects").status_code == 403
    assert client.get(f"/projects/{project_id}/drafts").status_code == 403
    assert (
        client.get(f"/projects/{project_id}/drafts/{policy_draft['id']}/download").status_code
        == 403
    )


def test_review_api_requires_login(client, db, policy_draft):
    """로그인 없이는 401 이다."""
    task_id = task_of(db, policy_draft["id"]).id
    client.cookies.clear()

    assert client.get("/reviews/queue").status_code == 401
    assert open_task(client, task_id).status_code == 401


def test_unknown_task_is_404(client, db, tenants):
    """없는 과제 ID 는 404 다."""
    login(client, REVIEWER_EMAIL)
    assert open_task(client, uuid.uuid4()).status_code == 404


# ---------------------------------------------------------------------------
# 큐·배정
# ---------------------------------------------------------------------------


def test_queue_shows_pending_task_with_org_context(client, db, policy_draft, tenants):
    """큐 항목에는 조직·프로젝트 이름과 `[확인 필요]` 통계가 함께 온다."""
    login(client, REVIEWER_EMAIL)
    response = client.get("/reviews/queue")
    assert response.status_code == 200, response.text

    items = response.json()
    assert len(items) == 1
    item = items[0]
    assert item["status"] == "pending"
    assert item["assigned_to_me"] is False

    draft = item["draft"]
    assert draft["id"] == policy_draft["id"]
    assert draft["kind"] == "policy"
    assert draft["version"] == 1
    # 심사원은 조직 API 를 못 쓰므로 이름이 응답에 실려 와야 한다.
    assert draft["project_name"] == tenants["project_a"].name
    assert draft["org_name"] == tenants["org_a"].name
    assert draft["stats"]["needs_review"] > 0


def test_opening_task_claims_it(client, db, policy_draft, tenants):
    """미배정 과제를 열면 그 심사원에게 배정된다."""
    login(client, REVIEWER_EMAIL)
    task_id = task_of(db, policy_draft["id"]).id

    response = open_task(client, task_id)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["assigned_to_me"] is True
    assert payload["reviewer_id"] == str(tenants["reviewer"].id)
    # 상세에는 편집 대상 본문이 통째로 실린다.
    assert len(payload["content_json"]["sections"]) >= 10

    assert task_of(db, policy_draft["id"]).reviewer_id == tenants["reviewer"].id
    assert client.get("/reviews/queue").json()[0]["assigned_to_me"] is True


def test_task_claimed_by_other_reviewer_is_forbidden(
    client, db, policy_draft, other_reviewer
):
    """남이 잡은 과제는 큐에서도 사라지고 직접 열어도 403 이다."""
    task_id = task_of(db, policy_draft["id"]).id

    login(client, REVIEWER_EMAIL)
    assert open_task(client, task_id).status_code == 200

    login(client, OTHER_REVIEWER_EMAIL)
    assert client.get("/reviews/queue").json() == []
    assert open_task(client, task_id).status_code == 403
    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 403


def test_operator_can_read_queue_but_not_decide(client, db, policy_draft):
    """운영자는 전 조직 열람 권한으로 큐를 보지만 결정은 하지 못한다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, "operator@example.com")

    assert len(client.get("/reviews/queue").json()) == 1
    detail = open_task(client, task_id)
    assert detail.status_code == 200
    # 열람만으로 배정되지 않는다.
    assert detail.json()["reviewer_id"] is None
    assert task_of(db, policy_draft["id"]).reviewer_id is None

    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 403
    assert client.post(f"/reviews/{task_id}/return", json={"comment": "보완"}).status_code == 403


# ---------------------------------------------------------------------------
# 편집 → DOCX 재생성
# ---------------------------------------------------------------------------


def test_editing_sow_row_updates_content_and_docx(client, db, sow_draft, storage, tenants):
    """행을 고치면 본문·`[확인 필요]` 통계·DOCX 가 모두 함께 갱신된다."""
    task_id = task_of(db, sow_draft.id).id
    login(client, REVIEWER_EMAIL)
    assert open_task(client, task_id).status_code == 200

    new_status = "정보보호 위원회를 분기 1회 개최한다."
    response = client.patch(
        f"/reviews/{task_id}/content",
        json={
            "rows": [
                {
                    "row_index": 0,
                    "fields": {
                        "operation_status": new_status,
                        "owner_dept": "정보보호팀",
                        "note": "심사원 확인 완료",
                    },
                }
            ]
        },
    )
    assert response.status_code == 200, response.text

    content = response.json()["content_json"]
    row = content["rows"][0]
    assert row["operation_status"] == new_status
    assert row["owner_dept"] == "정보보호팀"
    # 항목 코드·근거처럼 편집 대상이 아닌 칸은 그대로다.
    assert row["criterion_code"] == "1.1.1"
    assert row["related_refs"] == ["정보보호정책 v2.1 p.3"]
    # 채운 만큼 `[확인 필요]` 가 줄어든다.
    assert content["stats"]["needs_review"] == 2
    assert content["stats"]["needs_review_rows"] == 2

    draft = load_draft(db, sow_draft.id)
    assert draft.content_json["rows"][0]["operation_status"] == new_status
    # 편집과 동시에 DOCX 를 다시 만들어 같은 키에 올린다.
    assert draft.docx_s3_key == draft_docx_key(
        org_id=tenants["project_a"].org_id,
        project_id=draft.project_id,
        draft_id=draft.id,
        kind=DraftKind.SOW,
        version=1,
    )
    table = DocxDocument(io.BytesIO(storage.get_object(draft.docx_s3_key))).tables[0]
    assert table.rows[1].cells[2].text == new_status
    assert table.rows[1].cells[4].text == "정보보호팀"


def test_editing_policy_section_updates_body(client, db, policy_draft, storage):
    """정책 초안은 조항 본문 단위로 고친다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, REVIEWER_EMAIL)
    assert open_task(client, task_id).status_code == 200

    body = "본 정책은 2026년 1월 1일부터 시행한다."
    response = client.patch(
        f"/reviews/{task_id}/content",
        json={"sections": [{"section_index": 0, "body": body}]},
    )
    assert response.status_code == 200, response.text
    assert response.json()["content_json"]["sections"][0]["body"] == body

    draft = load_draft(db, policy_draft["id"])
    assert draft.content_json["sections"][0]["body"] == body
    texts = [
        paragraph.text
        for paragraph in DocxDocument(io.BytesIO(storage.get_object(draft.docx_s3_key))).paragraphs
    ]
    assert body in texts


def test_edit_is_audited_without_document_body(client, db, sow_draft, storage):
    """편집 이력이 감사 로그에 남는다. 본문은 남기지 않는다."""
    task_id = task_of(db, sow_draft.id).id
    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    client.patch(
        f"/reviews/{task_id}/content",
        json={"rows": [{"row_index": 1, "fields": {"owner_dept": "인프라팀"}}]},
    )

    db.expire_all()
    log = db.execute(
        select(AuditLog).where(AuditLog.action == "review.edit", AuditLog.target == str(task_id))
    ).scalar_one()
    assert log.meta_json["rows"] == [1]
    assert log.meta_json["kind"] == "sow"
    assert log.meta_json["needs_review"] == 3
    assert "인프라팀" not in str(log.meta_json)


def test_edit_rejects_wrong_shape(client, db, sow_draft, storage):
    """없는 행 번호와 초안 종류가 다른 편집은 400 이다."""
    task_id = task_of(db, sow_draft.id).id
    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)

    missing_row = client.patch(
        f"/reviews/{task_id}/content",
        json={"rows": [{"row_index": 99, "fields": {"note": "x"}}]},
    )
    assert missing_row.status_code == 400

    wrong_kind = client.patch(
        f"/reviews/{task_id}/content",
        json={"sections": [{"section_index": 0, "body": "x"}]},
    )
    assert wrong_kind.status_code == 400


# ---------------------------------------------------------------------------
# 승인 → 다운로드 개방
# ---------------------------------------------------------------------------


def test_approval_opens_download_for_org_admin(client, db, policy_draft, tenants, storage):
    """승인해야만 조직 관리자가 산출물을 내려받을 수 있다(데모 기준 D5)."""
    project_id = tenants["project_a"].id
    draft_id = policy_draft["id"]
    task_id = task_of(db, draft_id).id

    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    response = client.post(f"/reviews/{task_id}/approve", json={})
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["status"] == "approved"
    assert payload["decided_at"] is not None
    assert payload["draft"]["status"] == "approved"

    task = task_of(db, draft_id)
    assert task.status is ReviewTaskStatus.APPROVED
    assert load_draft(db, draft_id).status is DraftStatus.APPROVED

    login(client, "admin-a@example.com")
    detail = client.get(f"/projects/{project_id}/drafts/{draft_id}")
    assert detail.json()["downloadable"] is True
    download = client.get(f"/projects/{project_id}/drafts/{draft_id}/download")
    assert download.status_code == 200
    assert download.content[:2] == b"PK"


def test_approval_writes_audit_and_alert(client, db, policy_draft, tenants):
    """승인은 감사 로그와 조직 알림을 남긴다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 200

    db.expire_all()
    log = db.execute(
        select(AuditLog).where(
            AuditLog.action == "review.approve", AuditLog.target == str(task_id)
        )
    ).scalar_one()
    assert log.org_id == tenants["org_a"].id
    assert log.user_id == tenants["reviewer"].id
    assert log.meta_json["draft_id"] == policy_draft["id"]

    alert = db.execute(
        select(Alert).where(Alert.project_id == tenants["project_a"].id)
    ).scalar_one()
    # alerts.type 은 drift/due/defect 뿐이라 승인은 `due` 로 남긴다.
    assert alert.type is AlertType.DUE
    assert "승인 완료" in alert.message
    assert "다운로드" in alert.message


# ---------------------------------------------------------------------------
# 반려
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("body", [{}, {"comment": ""}, {"comment": "   "}])
def test_return_requires_comment(client, db, policy_draft, body):
    """반려에는 사유가 반드시 있어야 한다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)

    response = client.post(f"/reviews/{task_id}/return", json=body)
    assert response.status_code == 400
    assert response.json()["detail"] == "반려 사유를 입력해야 합니다"
    # 실패한 반려는 초안 상태를 건드리지 않는다.
    assert load_draft(db, policy_draft["id"]).status is DraftStatus.IN_REVIEW


def test_return_marks_draft_and_alerts_org(client, db, policy_draft, tenants, storage):
    """반려하면 초안이 되돌아가고 조직에 코멘트가 담긴 알림이 뜬다."""
    project_id = tenants["project_a"].id
    draft_id = policy_draft["id"]
    task_id = task_of(db, draft_id).id
    comment = "3조 보관 기간이 내부 규정과 다릅니다. 확인 후 다시 제출해 주세요."

    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    response = client.post(f"/reviews/{task_id}/return", json={"comment": comment})
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "returned"
    assert response.json()["comment"] == comment

    assert load_draft(db, draft_id).status is DraftStatus.RETURNED

    alert = db.execute(select(Alert).where(Alert.project_id == project_id)).scalar_one()
    assert alert.type is AlertType.DEFECT
    assert comment in alert.message

    db.expire_all()
    log = db.execute(
        select(AuditLog).where(AuditLog.action == "review.return", AuditLog.target == str(task_id))
    ).scalar_one()
    assert log.meta_json["comment"] == comment

    # 반려된 초안은 여전히 내려받을 수 없다.
    login(client, "admin-a@example.com")
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}/download").status_code == 403


def test_decided_task_cannot_be_decided_again(client, db, policy_draft, storage):
    """이미 결정된 과제는 다시 결정할 수 없고 편집도 막힌다."""
    task_id = task_of(db, policy_draft["id"]).id
    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 200

    assert client.post(f"/reviews/{task_id}/approve", json={}).status_code == 409
    assert client.post(f"/reviews/{task_id}/return", json={"comment": "번복"}).status_code == 409
    assert (
        client.patch(
            f"/reviews/{task_id}/content",
            json={"sections": [{"section_index": 0, "body": "x"}]},
        ).status_code
        == 409
    )


# ---------------------------------------------------------------------------
# 재제출
# ---------------------------------------------------------------------------


def test_resubmission_creates_new_version_and_task(client, db, policy_draft, tenants, storage):
    """반려 후 재생성하면 새 버전 초안이 새 과제로 큐에 다시 오른다."""
    project_id = tenants["project_a"].id
    task_id = task_of(db, policy_draft["id"]).id

    login(client, REVIEWER_EMAIL)
    open_task(client, task_id)
    client.post(f"/reviews/{task_id}/return", json={"comment": "보완이 필요합니다"})

    login(client, "admin-a@example.com")
    resubmitted = client.post(f"/projects/{project_id}/drafts", json={"kind": "policy"})
    assert resubmitted.status_code == 201, resubmitted.text
    assert resubmitted.json()["version"] == 2
    assert resubmitted.json()["status"] == "in_review"

    new_task = task_of(db, resubmitted.json()["id"])
    assert new_task.status is ReviewTaskStatus.PENDING
    assert new_task.reviewer_id is None

    login(client, REVIEWER_EMAIL)
    queue = client.get("/reviews/queue").json()
    # 대기 중인 새 과제가 먼저 오고, 처리한 이력이 뒤따른다.
    assert [item["status"] for item in queue] == ["pending", "returned"]
    assert queue[0]["draft"]["version"] == 2
