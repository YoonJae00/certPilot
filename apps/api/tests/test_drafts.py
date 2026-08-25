"""문서 코파일럿 테스트 (PRD §7 F4).

AC: 101행 생성, `[확인 필요]` 개수 리포트, DOCX 열림, 승인 전 다운로드 403.

실제 LLM 은 절대 부르지 않는다. `.env` 에 키가 있어도 결정적 Fake 로 고정한다.
"""

import io
import uuid

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.llm.provider import FakeProvider
from app.models import AuditLog, Draft, DraftStatus, Finding, FindingStatus
from app.services.draft_common import NEEDS_REVIEW, count_needs_review
from app.services.draft_docx import SOW_COLUMNS
from app.services.draft_sow import NO_ASSESSMENT_MESSAGE
from app.workers.assess import run_assessment
from tests.conftest import login
from tests.test_assess import RULE_FAIL_CODE, create_assessment, seed_project

TOTAL_CRITERIA = 101


@pytest.fixture(autouse=True)
def _no_celery(monkeypatch):
    """업로드가 실제 브로커를 건드리지 않게 한다."""
    monkeypatch.setattr("app.api.documents.enqueue_ingest", lambda document_id: None)


@pytest.fixture(autouse=True)
def _fake_llm(monkeypatch):
    """비용이 드는 실제 LLM 호출을 원천 차단한다."""
    monkeypatch.setattr("app.workers.assess.get_llm_provider", FakeProvider)


@pytest.fixture
def assessed(client, db, tenants, storage):
    """완료된 모의심사 1건이 있는 A조직 프로젝트(admin-a 로 로그인된 상태)."""
    project_id = seed_project(client, db, tenants)
    assessment_id = create_assessment(db, project_id)
    run_assessment(assessment_id, provider=FakeProvider())
    return {"project_id": project_id, "assessment_id": assessment_id}


def findings_by_code(db, assessment_id: uuid.UUID) -> dict[str, Finding]:
    """판정을 항목 코드로 인덱싱해 읽는다."""
    db.expire_all()
    rows = db.execute(select(Finding).where(Finding.assessment_id == assessment_id)).scalars()
    return {row.criterion_code: row for row in rows}


def create_draft(client, project_id, kind: str):
    """초안 생성 요청 헬퍼."""
    return client.post(f"/projects/{project_id}/drafts", json={"kind": kind})


def load_draft(db, draft_id) -> Draft:
    """DB 에서 초안을 다시 읽는다(API 와 다른 세션이다)."""
    db.expire_all()
    return db.execute(select(Draft).where(Draft.id == uuid.UUID(str(draft_id)))).scalar_one()


def approve(db, draft_id) -> None:
    """검수 API 없이 승인 상태만 직접 만든다(Task 9 전까지의 우회로)."""
    draft = load_draft(db, draft_id)
    draft.status = DraftStatus.APPROVED
    db.commit()


def row_of(content: dict, code: str) -> dict:
    """운영명세서에서 항목 코드로 행 1개를 찾는다."""
    for row in content["rows"]:
        if row["criterion_code"] == code:
            return row
    raise AssertionError(f"운영명세서에 항목이 없다: {code}")


# ---------------------------------------------------------------------------
# 운영명세서(sow)
# ---------------------------------------------------------------------------


def test_sow_draft_has_one_row_per_criterion(client, db, assessed):
    """101행이 만들어지고 in_review 로 생성된다."""
    response = create_draft(client, assessed["project_id"], "sow")
    assert response.status_code == 201, response.text

    payload = response.json()
    assert payload["kind"] == "sow"
    assert payload["version"] == 1
    # LLM 은 초안만 쓴다. 생성 즉시 검수 대기다.
    assert payload["status"] == "in_review"
    assert payload["downloadable"] is False

    content = payload["content_json"]
    assert len(content["rows"]) == TOTAL_CRITERIA
    assert content["stats"]["total"] == TOTAL_CRITERIA
    assert content["assessment_id"] == str(assessed["assessment_id"])

    codes = [row["criterion_code"] for row in content["rows"]]
    assert len(set(codes)) == TOTAL_CRITERIA
    # 사람이 보는 순서(2.2.1 → 2.10.1)로 정렬돼 있다.
    assert codes[0].startswith("1.")

    for row in content["rows"]:
        assert set(row) == {
            "criterion_code",
            "criterion_title",
            "section",
            "operation_status",
            "related_refs",
            "owner_dept",
            "note",
        }
        # 담당 부서는 시스템이 알 수 없다.
        assert row["owner_dept"] == NEEDS_REVIEW


def test_sow_needs_review_count_matches_rows(client, db, assessed):
    """`stats.needs_review` 가 실제 `[확인 필요]` 칸 수와 일치하고 판단불가 수 이상이다."""
    content = create_draft(client, assessed["project_id"], "sow").json()["content_json"]

    findings = findings_by_code(db, assessed["assessment_id"])
    unknown_codes = {
        code for code, item in findings.items() if item.status is FindingStatus.UNKNOWN
    }
    # 근거 없는 항목이 있어야 판단불가 초안 경로가 실제로 돈다.
    assert unknown_codes, "판단불가 판정이 하나도 없어 [확인 필요] 경로를 검증할 수 없다"

    recounted = count_needs_review(content["rows"])
    assert content["stats"]["needs_review"] == recounted
    assert content["stats"]["needs_review"] >= len(unknown_codes)

    marked_rows = [row for row in content["rows"] if NEEDS_REVIEW in row["operation_status"]]
    assert {row["criterion_code"] for row in marked_rows} >= unknown_codes
    assert content["stats"]["needs_review_rows"] == TOTAL_CRITERIA


def test_sow_unmet_row_quotes_predicted_defect(client, db, assessed):
    """미충족 항목 행에 예상 결함·개선 권고가 그대로 실린다."""
    content = create_draft(client, assessed["project_id"], "sow").json()["content_json"]
    finding = findings_by_code(db, assessed["assessment_id"])[RULE_FAIL_CODE]
    assert finding.status is FindingStatus.UNMET

    row = row_of(content, RULE_FAIL_CODE)
    assert row["operation_status"].startswith("현재 미이행.")
    assert finding.predicted_defect
    assert finding.predicted_defect in row["operation_status"]
    assert NEEDS_REVIEW not in row["operation_status"]
    # 규칙 fail 증적은 관련 증적 칸에 `source.check_id` 로 표시된다.
    assert "aws.iam.mfa_enabled" in row["related_refs"]


def test_sow_unknown_row_is_marked_needs_review(client, db, assessed):
    """판단불가 항목은 `[확인 필요]` + 판정하지 못한 사유로 채운다."""
    findings = findings_by_code(db, assessed["assessment_id"])
    unknown_code = next(
        code for code, item in findings.items() if item.status is FindingStatus.UNKNOWN
    )

    content = create_draft(client, assessed["project_id"], "sow").json()["content_json"]
    row = row_of(content, unknown_code)

    assert row["operation_status"].startswith(NEEDS_REVIEW)
    # 판정하지 못한 사유를 그대로 남긴다(없는 현황을 지어내지 않는다).
    assert "판단할 수 없다" in row["operation_status"] or "찾지 못" in row["operation_status"]
    assert row["related_refs"] == [NEEDS_REVIEW]


def test_sow_met_row_cites_document_and_page(client, db, assessed):
    """충족·부분충족 항목은 근거 문서명·쪽수를 관련 문서 칸에 표시한다."""
    findings = findings_by_code(db, assessed["assessment_id"])
    code = next(
        code
        for code, item in findings.items()
        if item.status in {FindingStatus.MET, FindingStatus.PARTIAL} and item.evidence_chunk_ids
    )

    content = create_draft(client, assessed["project_id"], "sow").json()["content_json"]
    row = row_of(content, code)

    assert row["related_refs"] != [NEEDS_REVIEW]
    # 파일명 대신 사람이 읽는 문서명이 들어간다(`01_...pdf` 가 아니다).
    assert not any(ref.endswith(".pdf") for ref in row["related_refs"])
    assert any("정보보호정책 v2.1" in ref for ref in row["related_refs"])
    # 운영 현황은 판정 근거 + 청크 발췌로 만든 서술형 초안이다.
    assert "근거 문서" in row["operation_status"]
    # 내부 식별자(chunk:c_…)는 사람이 읽는 문서에 남지 않는다.
    assert "chunk:c_" not in row["operation_status"]


def test_sow_without_assessment_is_400(client, db, tenants):
    """모의심사 없이 운영명세서를 만들 수 없다."""
    login(client, "admin-a@example.com")
    response = create_draft(client, tenants["project_a"].id, "sow")
    assert response.status_code == 400
    assert response.json()["detail"] == NO_ASSESSMENT_MESSAGE


def test_sow_version_increments(client, db, assessed):
    """같은 종류를 다시 만들면 버전이 올라간다."""
    first = create_draft(client, assessed["project_id"], "sow")
    second = create_draft(client, assessed["project_id"], "sow")
    assert first.json()["version"] == 1
    assert second.json()["version"] == 2

    listed = client.get(f"/projects/{assessed['project_id']}/drafts")
    assert listed.status_code == 200
    assert {item["version"] for item in listed.json()} == {1, 2}


# ---------------------------------------------------------------------------
# 정책 초안(policy)
# ---------------------------------------------------------------------------


def test_policy_draft_fills_template(client, db, tenants, storage):
    """정책 초안은 조항 10개 이상, 모르는 값은 `[확인 필요]` 로 남긴다."""
    login(client, "admin-a@example.com")
    project = tenants["project_a"]
    response = create_draft(client, project.id, "policy")
    assert response.status_code == 201, response.text

    payload = response.json()
    assert payload["kind"] == "policy"
    assert payload["status"] == "in_review"

    content = payload["content_json"]
    sections = content["sections"]
    assert len(sections) >= 10
    assert content["title"] == f"{project.name} 정보보호 정책"

    bodies = "\n".join(section["body"] for section in sections)
    # 회사명은 프로젝트명으로 채운다.
    assert project.name in bodies
    # 인증 종류는 프로젝트 설정에서 온다.
    assert project.cert_type.value in bodies
    # 서비스명·CISO·시행일은 프로젝트에 없는 값이라 사람이 채운다.
    assert NEEDS_REVIEW in bodies
    assert content["stats"]["needs_review"] == count_needs_review(sections)
    assert content["stats"]["needs_review"] > 0
    # 템플릿 플레이스홀더가 그대로 남으면 안 된다.
    assert "{{" not in bodies


def test_policy_draft_uses_project_scope(client, db, tenants, storage):
    """인증범위를 채운 프로젝트는 정책 본문에 그 문구가 들어간다."""
    project = tenants["project_a"]
    project.scope_text = "certpilot-saas 운영 조직과 AWS 서울 리전 인프라"
    db.commit()

    login(client, "admin-a@example.com")
    content = create_draft(client, project.id, "policy").json()["content_json"]
    bodies = "\n".join(section["body"] for section in content["sections"])
    assert project.scope_text in bodies


# ---------------------------------------------------------------------------
# DOCX 변환
# ---------------------------------------------------------------------------


def test_sow_docx_table_has_header_and_101_rows(client, db, assessed, storage):
    """S3 에 올라간 DOCX 를 열면 표가 머리글 1행 + 101행이다."""
    draft_id = create_draft(client, assessed["project_id"], "sow").json()["id"]
    draft = load_draft(db, draft_id)
    assert draft.docx_s3_key

    document = DocxDocument(io.BytesIO(storage.get_object(draft.docx_s3_key)))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == TOTAL_CRITERIA + 1
    assert len(table.columns) == len(SOW_COLUMNS)
    assert [cell.text for cell in table.rows[0].cells] == list(SOW_COLUMNS)

    first = table.rows[1].cells
    assert first[0].text.startswith("1.")
    assert first[4].text == NEEDS_REVIEW


def test_policy_docx_has_section_headings(client, db, tenants, storage):
    """정책 DOCX 는 제목 + 조항별 소제목·본문 문단으로 만들어진다."""
    login(client, "admin-a@example.com")
    response = create_draft(client, tenants["project_a"].id, "policy")
    content = response.json()["content_json"]
    draft = load_draft(db, response.json()["id"])

    document = DocxDocument(io.BytesIO(storage.get_object(draft.docx_s3_key)))
    texts = [paragraph.text for paragraph in document.paragraphs]
    for section in content["sections"]:
        assert section["heading"] in texts
    assert document.tables == []


# ---------------------------------------------------------------------------
# 승인 전 다운로드 차단 (데모 기준 D5)
# ---------------------------------------------------------------------------


def test_download_is_forbidden_before_approval(client, db, assessed, storage):
    """in_review 상태에서는 다운로드가 403 이다."""
    draft_id = create_draft(client, assessed["project_id"], "sow").json()["id"]

    response = client.get(f"/projects/{assessed['project_id']}/drafts/{draft_id}/download")
    assert response.status_code == 403
    assert response.json()["detail"] == "심사원 승인 후 다운로드할 수 있습니다"


def test_download_works_after_approval(client, db, assessed, storage):
    """승인되면 DOCX 바이트를 그대로 내려준다."""
    draft_id = create_draft(client, assessed["project_id"], "sow").json()["id"]
    approve(db, draft_id)

    detail = client.get(f"/projects/{assessed['project_id']}/drafts/{draft_id}")
    assert detail.json()["downloadable"] is True

    response = client.get(f"/projects/{assessed['project_id']}/drafts/{draft_id}/download")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "certpilot-sow-v1.docx" in response.headers["content-disposition"]
    # DOCX 는 zip 컨테이너다. python-docx 로 그대로 열린다.
    assert response.content[:2] == b"PK"
    assert len(DocxDocument(io.BytesIO(response.content)).tables[0].rows) == TOTAL_CRITERIA + 1


def test_create_and_download_are_audited(client, db, assessed, storage):
    """생성·다운로드가 감사 로그에 남는다(문서 본문은 남기지 않는다)."""
    draft_id = create_draft(client, assessed["project_id"], "sow").json()["id"]
    approve(db, draft_id)
    assert client.get(
        f"/projects/{assessed['project_id']}/drafts/{draft_id}/download"
    ).status_code == 200

    db.expire_all()
    created = db.execute(
        select(AuditLog).where(AuditLog.action == "draft.create", AuditLog.target == draft_id)
    ).scalar_one()
    assert created.meta_json["kind"] == "sow"
    assert created.meta_json["stats"]["total"] == TOTAL_CRITERIA
    assert "rows" not in created.meta_json

    downloaded = db.execute(
        select(AuditLog).where(AuditLog.action == "draft.download", AuditLog.target == draft_id)
    ).scalar_one()
    assert downloaded.meta_json["bytes"] > 0


# ---------------------------------------------------------------------------
# 권한·테넌트 격리
# ---------------------------------------------------------------------------


def test_other_org_cannot_touch_drafts(client, db, assessed, storage):
    """다른 조직은 존재 여부조차 알 수 없다(404)."""
    project_id = assessed["project_id"]
    draft_id = create_draft(client, project_id, "sow").json()["id"]

    login(client, "admin-b@example.com")
    assert create_draft(client, project_id, "sow").status_code == 404
    assert client.get(f"/projects/{project_id}/drafts").status_code == 404
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}").status_code == 404
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}/download").status_code == 404


def test_reviewer_cannot_use_org_scoped_draft_api(client, db, assessed, storage):
    """심사원은 조직 스코프 API 로 초안에 접근할 수 없다(403)."""
    project_id = assessed["project_id"]
    draft_id = create_draft(client, project_id, "sow").json()["id"]
    approve(db, draft_id)

    login(client, "reviewer@example.com")
    assert create_draft(client, project_id, "sow").status_code == 403
    assert client.get(f"/projects/{project_id}/drafts").status_code == 403
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}").status_code == 403
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}/download").status_code == 403


def test_member_can_read_but_not_create(client, db, assessed, storage):
    """조직 팀원은 열람만 할 수 있다."""
    project_id = assessed["project_id"]
    draft_id = create_draft(client, project_id, "sow").json()["id"]

    login(client, "member-a@example.com")
    assert create_draft(client, project_id, "policy").status_code == 403
    assert client.get(f"/projects/{project_id}/drafts").status_code == 200
    assert client.get(f"/projects/{project_id}/drafts/{draft_id}").status_code == 200


def test_requires_login(client, db, tenants):
    """로그인 없이는 401 이다."""
    project_id = tenants["project_a"].id
    assert create_draft(client, project_id, "sow").status_code == 401
    assert client.get(f"/projects/{project_id}/drafts").status_code == 401


def test_unknown_draft_id_is_404(client, db, assessed, storage):
    """없는 초안 ID 는 404 다."""
    project_id = assessed["project_id"]
    missing = uuid.uuid4()
    assert client.get(f"/projects/{project_id}/drafts/{missing}").status_code == 404
    assert client.get(f"/projects/{project_id}/drafts/{missing}/download").status_code == 404
